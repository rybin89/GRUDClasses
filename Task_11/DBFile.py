from docx import Document  # pip install python-docx
import os


class DBFile:
    '''
    класс для хранения информации о файле
    Проверить есть ли файле
        усли нет то создать
    Проверить есть ли таблица
        если нет то создать
    Читать файл
        получить все строки таблицы в виде списка словарей
        events = [{"id": 1, "title": "Встреча", "date": "2024-02-15", "time": "14:00","description": "С коллегами"}]
    Добавить строку в таблицу
        добавить строку в таблицу
            добавить в конец
    Attributes:
        doc (Document): объект для работы с файлом
        table (Table): таблица в файле
    Returns:
        none
    '''

    def __init__(self, path, headers):# путь к файлу, список заголовков таблицы
        self.path = path# создаем объект для работы с файлом
        self.headers = headers # список заголовков таблицы

        self._ensure_table() # проверяем наличие таблицы

    def create_file(self):# создание файла
        doc = Document()  # Создаем новый документ
        doc.add_heading("Календарь событий", 0) # добавляем заголовок в файл
        self._create_table(doc) # создаем таблицу в файле
        doc.save(self.path) # сохраняем файл
        return f'Создан новый файл {self.path}' # возвращаем сообщение о создании файла

    def _create_table(self, doc):# создание таблицы
        table = doc.add_table(rows=1, cols=len(self.headers)) # создаем таблицу в файле
        table.style = 'Table Grid'  # задаем стиль таблицы
        header_cell = table.rows[0].cells  # получаем первую строку таблицы
        for i, header in enumerate(self.headers):  # перебираем заголовки
            header_cell[i].text = header  # добавляем заголовок в таблицу

    def _ensure_table(self):# проверка наличия таблицы
        '''
        Description:
            Проверяет наличие таблицы в файле
        Returns:
            none
        '''
        if not os.path.exists(self.path):  # Проверяем, существует ли файл
            self.create_file()
        else:
            doc = Document(self.path)  # Открываем существующий документ
            if len(doc.tables) == 0:     # Проверяем, есть ли таблицы
                self._create_table(doc)
                doc.save(self.path)      # Сохраняем изменения
    def read(self):# чтение файла
        '''
        Чтение файла

        :return:
        список словарей
        '''
        list = []
        doc = Document(self.path)
        if len(doc.tables) == 0:
            return 'файл пуст'
        table = doc.tables[0] # получаем первую таблицу

        for i in range(1, len(table.rows)): # перебираем все строки таблицы начиная с 1
            row = table.rows[i] # получаем строку
            row_dict = {} # создаем словарь для строки
            for j in range(len(row.cells)): # перебираем все ячейки строки
                if j < len(self.headers): # проверяем, что индекс не выходит за границы списка заголовков
                    row_dict[self.headers[j]] = row.cells[j].text # добавляем ячейку в словарь
            list.append(row_dict) # добавляем словарь в список
        return list
    def add(self, row):# добавление строки в таблицу
        '''
        Добавление строки в таблицу
        возвращает сообщение об успешном добавлении
        :param row: dict
        :return:
        строка добавлена
        '''
        doc = Document(self.path)
        table = doc.tables[0]
        new_row = table.add_row().cells # добавляем новую строку
        for i in range(len(self.headers)): # перебираем все заголовки

            new_row[i].text = str(row[self.headers[i]] )# добавляем значение из словаря в ячейку
        doc.save(self.path) # сохраняем файл
        return 'строка добавлена'
if __name__ == "__main__":
    file_path = 'test.docx'
    headers = [ 'id',
        'title',
        'date',
        'time',
        'description']
    doc = DBFile(file_path, headers)
    print(doc.read())
    doc.add({
      'id': 1,
        'title': 'Встреча',
        'date': '2024-02-15',
        'time': '14:00',
        'description': 'С коллегами'
    })